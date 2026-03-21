"""
backend/agents/resume_agent.py
================================
Tailors resume via LLM → LaTeX → PDF.  Parses the base .docx for structure,
sends to LLM for FAANG-style rewriting, renders via LaTeX template, compiles
with tectonic.  No .docx output — only .tex + .pdf.
"""

from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.text.paragraph import Paragraph
from loguru import logger
from slugify import slugify

from backend.db.client import db
from backend.agents.latex_resume_agent import render_latex_resume, compile_pdf
from backend.prompts.resume_prompt import (
    RESUME_TAILOR_SYSTEM_PROMPT,
    build_resume_tailoring_user_message,
)
from backend.utils.llm_client import call_llm_sync, parse_json_response

PROFILE_PATH = Path("data/candidate_profile.json")
RESUME_BASE_PATH = Path("data/resumes/resume_base.docx")
TAILORED_DIR = Path("data/resumes/tailored")
REPO_ROOT = Path(__file__).resolve().parents[2]


@dataclass
class ExperienceBlock:
    header_paragraphs: list[Paragraph] = field(default_factory=list)
    bullet_paragraphs: list[Paragraph] = field(default_factory=list)


@dataclass
class ProjectBlock:
    title_paragraph: Paragraph
    tech_paragraph: Optional[Paragraph] = None
    bullet_paragraphs: list[Paragraph] = field(default_factory=list)
    orphan_paragraphs: list[Paragraph] = field(default_factory=list)


@dataclass
class ResumeStructure:
    summary_paragraphs: list[Paragraph] = field(default_factory=list)
    skills_paragraphs: list[Paragraph] = field(default_factory=list)
    experience_blocks: list[ExperienceBlock] = field(default_factory=list)
    project_blocks: list[ProjectBlock] = field(default_factory=list)


# ── Main entry ────────────────────────────────────────────────────


def run_resume_agent(job: dict) -> dict:
    if not RESUME_BASE_PATH.exists():
        raise FileNotFoundError(
            f"Base resume not found at {RESUME_BASE_PATH}. "
            "Place resume_base.docx under data/resumes/."
        )
    if not PROFILE_PATH.exists():
        raise FileNotFoundError(f"Candidate profile not found at {PROFILE_PATH}")

    profile = json.loads(PROFILE_PATH.read_text(encoding="utf-8"))
    job_title = (job.get("title") or "").strip()
    company = (job.get("company") or "").strip()
    description = (job.get("description") or "").strip()[:1500]

    logger.info(f"Resume agent: tailoring for {job_title} @ {company}")

    base_doc = Document(str(RESUME_BASE_PATH))
    structure = _parse_document_structure(base_doc)
    structured_for_llm = _structure_for_llm(structure)

    profile_excerpt = {
        "skills": profile.get("skills", {}),
        "work_experience": (profile.get("work_experience", []) or [])[:6],
        "projects": (profile.get("projects", []) or [])[:3],
        "education": profile.get("education", {}),
        "target_roles": (profile.get("target_roles", []) or [])[:5],
    }

    resume_text = "\n".join(
        (p.text or "").strip() for p in base_doc.paragraphs if (p.text or "").strip()
    )[:3500]

    user_msg = build_resume_tailoring_user_message(
        job_title=job_title,
        company=company,
        job_description_first_1500=description,
        resume_full_text=resume_text,
        structured_resume={"profile": profile_excerpt, **structured_for_llm},
    )

    raw = call_llm_sync(
        messages=[{"role": "user", "content": user_msg}],
        system=RESUME_TAILOR_SYSTEM_PROMPT,
        max_tokens=8000,
        temperature=0.2,
        expect_json=True,
    )
    data = parse_json_response(raw)
    if not isinstance(data, dict):
        raise ValueError("LLM returned non-object JSON")

    # ── LaTeX → PDF (single output file) ──────────────────────────
    TAILORED_DIR.mkdir(parents=True, exist_ok=True)
    company_slug = slugify(company)[:60] or "company"
    role_slug = slugify(job_title)[:60] or "role"
    basename = f"{company_slug}_{role_slug}"

    exp_meta = _extract_experience_meta(structure.experience_blocks)
    proj_meta = _extract_project_meta(structure.project_blocks)
    tex_content = render_latex_resume(data, exp_meta, proj_meta)
    tex_path, pdf_path = compile_pdf(tex_content, TAILORED_DIR, basename)

    pdf_path_str = (
        str(pdf_path.resolve().relative_to(REPO_ROOT.resolve())) if pdf_path else None
    )
    file_path_str = pdf_path_str or str(tex_path.resolve().relative_to(REPO_ROOT.resolve()))

    resume_id = _record_in_db(
        job_id=job.get("id"),
        version_name=f"tailored_{basename}",
        file_path=file_path_str,
        diff_summary=str(data.get("diff_summary") or ""),
        keywords_added=data.get("keywords_added") or [],
    )

    return {
        "resume_id": resume_id,
        "file_path": file_path_str,
        "pdf_path": pdf_path_str,
        "diff_summary": data.get("diff_summary", ""),
        "keywords_added": data.get("keywords_added", []),
        "skills_line": data.get("skills_line", ""),
        "jd_keywords": data.get("jd_keywords", []),
        "ats_score_estimate": data.get("ats_score_estimate"),
    }


# ── Parsing (reads base .docx for LLM prompt) ────────────────────


def _is_heading(p: Paragraph) -> bool:
    if not p.style or not p.style.name:
        return False
    return "heading" in p.style.name.lower()


def _is_list_paragraph(p: Paragraph) -> bool:
    st = (p.style.name or "").lower() if p.style else ""
    if "list" in st or "bullet" in st:
        return True
    try:
        return bool(p._p.xpath(".//w:numPr"))
    except Exception:
        return False


def _parse_document_structure(doc: Document) -> ResumeStructure:
    paragraphs = list(doc.paragraphs)
    struct = ResumeStructure()
    state = "none"
    cur_headers: list[Paragraph] = []
    cur_bullets: list[Paragraph] = []
    cur_proj: Optional[ProjectBlock] = None

    def flush_exp() -> None:
        nonlocal cur_headers, cur_bullets
        if cur_headers or cur_bullets:
            struct.experience_blocks.append(
                ExperienceBlock(
                    header_paragraphs=cur_headers.copy(),
                    bullet_paragraphs=cur_bullets.copy(),
                )
            )
        cur_headers = []
        cur_bullets = []

    i = 0
    n = len(paragraphs)
    while i < n:
        p = paragraphs[i]
        text = (p.text or "").strip()
        tl = text.lower()

        if _is_heading(p):
            if "professional summary" in tl or tl == "summary":
                state = "summary"
                i += 1
                while i < n and not _is_heading(paragraphs[i]):
                    if (paragraphs[i].text or "").strip():
                        struct.summary_paragraphs.append(paragraphs[i])
                    i += 1
                continue

            if "skill" in tl:
                state = "skills"
                i += 1
                while i < n and not _is_heading(paragraphs[i]):
                    if (paragraphs[i].text or "").strip():
                        struct.skills_paragraphs.append(paragraphs[i])
                    i += 1
                continue

            if "experience" in tl and "project" not in tl:
                flush_exp()
                state = "experience"
                i += 1
                continue

            if "project" in tl and "experience" not in tl:
                flush_exp()
                state = "projects"
                cur_proj = None
                i += 1
                continue

            if "education" in tl:
                if state == "experience":
                    flush_exp()
                state = "none"
                i += 1
                continue

        if state == "experience":
            if _is_heading(p):
                tlh = tl
                if "project" in tlh:
                    flush_exp()
                    state = "projects"
                    cur_proj = None
                    i += 1
                    continue
                if "education" in tlh:
                    flush_exp()
                    state = "none"
                    i += 1
                    continue
            if _is_list_paragraph(p):
                cur_bullets.append(p)
            else:
                if cur_bullets:
                    flush_exp()
                if text:
                    cur_headers.append(p)
            i += 1
            continue

        if state == "projects":
            if _is_heading(p) and "education" in tl:
                break
            st = (p.style.name or "").lower() if p.style else ""
            if st.startswith("heading 2"):
                cur_proj = ProjectBlock(title_paragraph=p)
                struct.project_blocks.append(cur_proj)
                i += 1
                if i < n and not _is_list_paragraph(paragraphs[i]) and not _is_heading(paragraphs[i]):
                    tnext = (paragraphs[i].text or "").strip()
                    if tnext:
                        cur_proj.tech_paragraph = paragraphs[i]
                    i += 1
                continue
            if cur_proj is None:
                i += 1
                continue
            if cur_proj.tech_paragraph is None and not _is_list_paragraph(p) and text:
                cur_proj.tech_paragraph = p
                i += 1
                continue
            if _is_list_paragraph(p):
                cur_proj.bullet_paragraphs.append(p)
                i += 1
                continue
            if text:
                cur_proj.orphan_paragraphs.append(p)
            i += 1
            continue

        i += 1

    if state == "experience":
        flush_exp()

    return struct


def _structure_for_llm(struct: ResumeStructure) -> dict[str, Any]:
    experience = []
    for block in struct.experience_blocks:
        role_line = (block.header_paragraphs[0].text or "").strip() if block.header_paragraphs else ""
        company_line = (block.header_paragraphs[1].text or "").strip() if len(block.header_paragraphs) > 1 else ""
        bullets_orig = [(bp.text or "").strip() for bp in block.bullet_paragraphs]
        experience.append({
            "role_line": role_line,
            "company_line": company_line,
            "bullets_original": bullets_orig,
        })
    projects = []
    for pb in struct.project_blocks:
        projects.append({
            "name": (pb.title_paragraph.text or "").strip(),
            "tech_original": (pb.tech_paragraph.text or "").strip() if pb.tech_paragraph else "",
            "bullets_original": [(b.text or "").strip() for b in pb.bullet_paragraphs],
        })
    return {"experience": experience, "projects": projects}


# ── Metadata extraction (for LaTeX renderer) ─────────────────────


def _extract_experience_meta(blocks: list[ExperienceBlock]) -> list[dict[str, str]]:
    meta = []
    for b in blocks:
        role_text = (b.header_paragraphs[0].text or "") if b.header_paragraphs else ""
        company_text = (b.header_paragraphs[1].text or "") if len(b.header_paragraphs) > 1 else ""
        rp = role_text.split("\t")
        cp = company_text.split("\t")
        meta.append({
            "role": rp[0].strip() if rp else "",
            "dates": rp[1].strip() if len(rp) > 1 else "",
            "company": cp[0].strip() if cp else "",
            "location": cp[1].strip() if len(cp) > 1 else "",
        })
    return meta


def _extract_project_meta(blocks: list[ProjectBlock]) -> list[dict[str, str]]:
    meta = []
    for pb in blocks:
        title = (pb.title_paragraph.text or "").strip()
        parts = title.split("\t")
        tech = (pb.tech_paragraph.text or "").strip() if pb.tech_paragraph else ""
        meta.append({
            "name": parts[0].strip() if parts else "",
            "tech": tech,
        })
    return meta


# ── DB ────────────────────────────────────────────────────────────


def _record_in_db(
    *,
    job_id: Optional[str],
    version_name: str,
    file_path: str,
    diff_summary: str,
    keywords_added: list,
) -> Optional[str]:
    if not job_id:
        logger.warning("No job_id — skipping DB write for resume version")
        return None
    try:
        client = db()
        r = client.table("resumes").insert({
            "job_id": job_id,
            "version_name": version_name,
            "resume_type": "tailored",
            "file_path": file_path,
            "diff_summary": diff_summary,
            "keywords_added": [str(k) for k in keywords_added if k],
        }).execute()
        if r.data and len(r.data) > 0 and r.data[0].get("id"):
            return str(r.data[0]["id"])
    except Exception as e:
        logger.error(f"DB write failed for resume version: {e}")
    return None
